#!/usr/bin/env python3
"""Build a Word simulation report from a sweep_runner.py manifest (or any
case_id-keyed CSV of extracted results).

Every number in the report comes directly from the manifest columns; this
tool never fabricates or interpolates values beyond what the file contains.
"""

from __future__ import annotations

import argparse
import csv
import sys
from pathlib import Path
from typing import Any


class ReportError(RuntimeError):
    """Expected, user-facing error for this tool."""


def load_manifest(path: Path) -> list[dict[str, str]]:
    if not path.exists():
        raise ReportError(f"manifest not found: {path}")
    with path.open("r", encoding="utf-8", newline="") as fh:
        reader = csv.DictReader(fh)
        rows = list(reader)
    if not rows:
        raise ReportError("manifest contained no rows")
    return rows


def numeric_columns(rows: list[dict[str, str]]) -> list[str]:
    exclude = {"case_id", "status", "error", "stderr", "raw_last_line"}
    candidates = [k for k in rows[0].keys() if k not in exclude]
    numeric = []
    for col in candidates:
        try:
            float(rows[0][col])
            numeric.append(col)
        except (TypeError, ValueError):
            continue
    return numeric


def build_docx(rows: list[dict[str, str]], out_path: Path, title: str, template: str) -> None:
    try:
        import docx
        from docx.shared import Inches
    except ImportError:
        raise ReportError("python-docx is required; install with 'pip install python-docx'")

    doc = docx.Document()
    doc.add_heading(title, level=1)

    ok_rows = [r for r in rows if r.get("status", "ok") == "ok"]
    failed_rows = [r for r in rows if r.get("status", "ok") != "ok"]

    doc.add_paragraph(
        f"Cases reported: {len(rows)} total, {len(ok_rows)} succeeded, {len(failed_rows)} failed."
    )
    if failed_rows:
        doc.add_heading("Failed cases", level=2)
        for row in failed_rows:
            doc.add_paragraph(f"{row.get('case_id')}: {row.get('error', 'unknown error')}", style="List Bullet")

    if not ok_rows:
        doc.save(out_path)
        return

    doc.add_heading("Results table", level=2)
    columns = list(ok_rows[0].keys())
    columns = [c for c in columns if c not in ("status",)]
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Light Grid Accent 1"
    header_cells = table.rows[0].cells
    for i, col in enumerate(columns):
        header_cells[i].text = col
    for row in ok_rows:
        cells = table.add_row().cells
        for i, col in enumerate(columns):
            cells[i].text = str(row.get(col, ""))

    metric_cols = numeric_columns(ok_rows)
    if template == "comparison" and metric_cols and len(ok_rows) > 1:
        doc.add_heading("Comparison charts", level=2)
        try:
            import matplotlib

            matplotlib.use("Agg")
            import matplotlib.pyplot as plt
        except ImportError:
            doc.add_paragraph("(matplotlib not installed; charts skipped)")
        else:
            case_ids = [r["case_id"] for r in ok_rows]
            for metric in metric_cols:
                values = [float(r[metric]) for r in ok_rows]
                fig, ax = plt.subplots(figsize=(6, 3))
                ax.bar(case_ids, values)
                ax.set_ylabel(metric)
                ax.set_title(metric)
                fig.tight_layout()
                img_path = out_path.parent / f"_chart_{metric}.png"
                fig.savefig(img_path, dpi=150)
                plt.close(fig)
                doc.add_picture(str(img_path), width=Inches(5.5))
                img_path.unlink(missing_ok=True)

    doc.save(out_path)


def cmd_build(args: argparse.Namespace) -> int:
    manifest_path = Path(args.manifest)
    rows = load_manifest(manifest_path)
    out_path = Path(args.out)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    title = args.title or f"Simulation report — {manifest_path.stem}"
    build_docx(rows, out_path, title=title, template=args.template)

    print(f"Report written to {out_path} ({len(rows)} case(s) from {manifest_path})")
    return 0


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Build a Word report from a sweep manifest CSV.")
    sub = parser.add_subparsers(dest="command", required=True)

    build = sub.add_parser("build", help="Render a .docx report from a manifest CSV")
    build.add_argument("--manifest", required=True, help="Path to manifest.csv")
    build.add_argument("--out", required=True, help="Output .docx path")
    build.add_argument("--title", default=None, help="Report title (default derived from manifest filename)")
    build.add_argument(
        "--template",
        choices=["basic", "comparison"],
        default="basic",
        help="'basic' for a single case, 'comparison' to add per-metric bar charts across cases",
    )
    build.set_defaults(func=cmd_build)

    return parser


def main(argv: list[str] | None = None) -> int:
    parser = build_parser()
    args = parser.parse_args(argv)
    try:
        return args.func(args)
    except ReportError as exc:
        print(f"error: {exc}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    sys.exit(main())
